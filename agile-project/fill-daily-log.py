"""Generate / refresh the Project Daily Log for Smart School Management System."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).parent / "Daily Log - Smart School Management System.docx"

TODAY = date.today().strftime("%d %B %Y")

# (Date, Person Responsible, Action)
LOG_ENTRIES: list[tuple[str, str, str]] = [
    (
        "17 June 2026",
        "G.G. Kusumsiri (Principal)",
        "Issued formal project mandate requesting a cloud-based school management system pilot at CP/KOT/Delta Gemunupura College, Pussellawa.",
    ),
    (
        "17 June 2026",
        "G.G.M.P. Kusumsiri (Start-Up Manager)",
        "Drafted Project Brief (CS11) under PRINCE2 Starting Up a Project; distributed to project team for review.",
    ),
    (
        "17 June 2026",
        "G.G.M.P. Kusumsiri (Start-Up Manager / Developer Team Leader)",
        "Completed Project Plan (DGC-DBMS-PP-001) from PRINCE2 template using fill-project-plan.py.",
    ),
    (
        "17 June 2026",
        "U.G.C. Nayanathara (Project Manager)",
        "Confirmed Phase 1 scope with team: students, staff, inventory, examinations, user management, PWA, and Android APK.",
    ),
    (
        "17 June 2026",
        "G.M.D. Fernando (Risk Manager)",
        "Documented initial risk register (R01–R10) in Project Plan section 10.",
    ),
    (
        "17 June 2026",
        "Developer team (Leader: G.G.M.P. Kusumsiri)",
        "Verified existing application baseline: Next.js 16, Firebase Auth/Firestore, role-based dashboard, and core CRUD modules.",
    ),
    (
        "18 June 2026",
        "G.G.M.P. Kusumsiri (Start-Up Manager)",
        "Finalised Project Plan formatting (org structure, MoSCoW scope, deliverables, controls, and risk tables) using format-project-plan.py.",
    ),
    (
        "18 June 2026",
        "G.G.M.P. Kusumsiri (Developer Team Leader)",
        "Reviewed Phase 1 modules against mandate: student/staff management, examinations, inventory, user management, PWA, and Android TWA workflow.",
    ),
    (
        "18 June 2026",
        "T.H. Samaranayake (Quality Manager)",
        "Reviewed quality baseline: Vitest unit tests (access control, exam grading, Quick PIN), GitHub Actions CI, and Firestore security rules.",
    ),
    (
        "18 June 2026",
        "N.D. Ranasinghagei (Scheduling Manager)",
        "Confirmed two-week sprint cycle and Phase 1 pilot go-live target of August 2026.",
    ),
    (
        "19 June 2026",
        "G.G.M.P. Kusumsiri (Start-Up Manager)",
        "Created and issued Project Daily Log (DGC-DBMS-DL-001) for ongoing activity recording.",
    ),
    (
        "19 June 2026",
        "G.G.M.P. Kusumsiri (Start-Up Manager)",
        "Conducted full repository analysis; confirmed Timetable module deferred to Phase 2; documented pilot readiness status for team review.",
    ),
    (
        "19 June 2026",
        "U.G.C. Nayanathara (Project Manager)",
        "Prepared fortnightly demo agenda for Principal: dashboard, examination results entry, parent portal, and Android install guide.",
    ),
]

BLANK_ROWS = 22

COVER_ROWS = [
    ("Project:", "Smart School Management System"),
    ("Release:", "Phase 1: Pilot"),
    ("Date:", TODAY),
    ("Document type:", "Project Daily Log"),
    ("Author:", "G.G.M.P. Kusumsiri"),
    ("Owner:", "G.G.M.P. Kusumsiri (Start-Up Manager)"),
    (
        "Client:",
        "G.G. Kusumsiri,\nPrincipal (SLPS-1),\nCP/KOT/Delta Gemunupura College,\nPussellawa",
    ),
    ("Document Ref:", "DGC-DBMS-DL-001"),
    ("Version No:", "1.0"),
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def apply_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_document() -> Document:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT DAILY LOG")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Daily activity log")
    sub_run.bold = True
    sub_run.font.size = Pt(12)

    intro = doc.add_paragraph(
        "Record project activities below. One row per action. "
        "The Start-Up Manager (G.G.M.P. Kusumsiri) maintains this log; "
        "all team members may add entries after stand-up or at end of day."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Generated: {TODAY}  |  DGC-DBMS-DL-001 v1.0")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    cover = doc.add_table(rows=len(COVER_ROWS), cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(cover)
    for i, (label, value) in enumerate(COVER_ROWS):
        set_cell_text(cover.rows[i].cells[0], label, bold=True, size=10)
        set_cell_text(cover.rows[i].cells[1], value, size=10)
        if i == 0:
            set_cell_shading(cover.rows[i].cells[0], "D9E2F3")
            set_cell_shading(cover.rows[i].cells[1], "D9E2F3")

    doc.add_paragraph()

    log = doc.add_table(rows=1 + len(LOG_ENTRIES) + BLANK_ROWS, cols=3)
    log.alignment = WD_TABLE_ALIGNMENT.CENTER
    apply_table_borders(log)

    headers = ["Date", "Person Responsible", "Action"]
    for i, header in enumerate(headers):
        set_cell_text(log.rows[0].cells[i], header, bold=True, size=10)
        set_cell_shading(log.rows[0].cells[i], "1F4E79")
        for run in log.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (dt, person, action) in enumerate(LOG_ENTRIES, start=1):
        set_cell_text(log.rows[row_idx].cells[0], dt, size=10)
        set_cell_text(log.rows[row_idx].cells[1], person, size=10)
        set_cell_text(log.rows[row_idx].cells[2], action, size=10)

    for row_idx in range(1 + len(LOG_ENTRIES), 1 + len(LOG_ENTRIES) + BLANK_ROWS):
        for col in range(3):
            set_cell_text(log.rows[row_idx].cells[col], "", size=10)

    for row in log.rows:
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(4.4)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
