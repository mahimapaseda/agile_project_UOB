"""Reformat Project Plan docx: tables, bullets, aligned sections."""

from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

DOC_PATH = Path(__file__).parent / "Project Plan - Delta Gemunupura DBMS.docx"

TODAY = date.today().strftime("%d %B %Y")

RISKS = [
    ("R01", "Internet dependency — no offline database", "High",
     "Plan PWA offline page; communicate network requirement to users", "G.M.D. Fernando"),
    ("R02", "Firebase cost overrun at scale", "Medium",
     "Monitor usage; optimise queries; publish composite indexes", "G.M.D. Fernando"),
    ("R03", "Teachers resist adopting new system", "Medium",
     "Training session; Quick PIN login; parallel paper records during pilot", "G.M.D. Fernando"),
    ("R04", "Incorrect RBAC exposes student data", "High",
     "Design Firestore rules and access-control tests; verify VP view-only", "G.M.D. Fernando"),
    ("R05", "iOS users expect native App Store app", "Low",
     "Provide iOS install guide; document exclusion in Project Brief", "G.M.D. Fernando"),
    ("R06", "Android Play Protect warnings (WebAPK)", "Medium",
     "Plan official signed Android APK; document install steps for IT", "G.M.D. Fernando"),
    ("R07", "Service account / PIN pepper misconfigured for production", "High",
     "Use production environment setup checklist; planned health-check endpoint", "G.M.D. Fernando"),
    ("R08", "Timetable module not ready for pilot", "Medium",
     "Defer to Phase 2; agree scope with Principal", "G.M.D. Fernando"),
    ("R09", "Student login linking errors (wrong admission number)", "Medium",
     "Plan bulk provision process; admin verification before go-live", "G.M.D. Fernando"),
    ("R10", "Academic calendar pressure", "Medium",
     "Prioritise exam module in sprint plan; agile reprioritisation each sprint", "G.M.D. Fernando"),
]

MOSCOW_ROWS = [
    ("Must have", [
        "Student Management (CRUD, search, CSV import, PDF export)",
        "Staff Management (CRUD, Google Form CSV/XLSX import)",
        "Examination Management (create exams, enter results, auto-grading, ranks, report export)",
        "Inventory Management (asset register CRUD)",
        "User Management (principal, technical officer, VP, teacher, student, parent logins)",
        "Web PWA with install support",
        "Android official APK (Trusted Web Activity)",
        "Real-time performance analytics on data entry",
        "Parent portal (view child's progress and exam results)",
    ]),
    ("Should have", [
        "Quick PIN login for mobile users",
        "Biometric unlock on installed PWA",
        "WhatsApp integration for sharing student credentials",
        "Examination information reports for leadership",
    ]),
    ("Could have", [
        "Timetable Management module",
        "Teacher performance percentage dashboards",
        "Bulk SMS notifications to parents",
    ]),
    ("Won't have (Phase 1)", [
        "Native iOS App Store application (Safari PWA install only)",
        "Full offline database synchronisation",
        "Integration with government EMIS or national exam systems",
        "Fee/billing or payroll modules",
    ]),
]

DELIVERABLES = [
    "Production web application (PWA) for school use",
    "Android APK (Trusted Web Activity) with install guide",
    "Firebase project with Authentication, Firestore, security rules, and composite indexes",
    "Source code repository with CI workflow",
    "User documentation: system guide, Android install guide, and iOS install instructions",
    "Sample CSV templates for student, staff, exam results, and inventory import",
    "Admin setup utilities for creating the first principal account and provisioning student logins",
    "Project documentation: Project Mandate, Project Brief, and this Project Plan",
    "Role-based dashboards and module screens for all mandated user types",
    "PDF export for student/staff profiles and examination reports",
]

ORG_ROWS = [
    ("Executive / Client", "G.G. Kusumsiri", "Principal — project mandate, scope approval, and final acceptance"),
    ("Senior User", "Vice Principal", "School operations representative"),
    ("Senior Supplier", "U.G.C. Nayanathara (also Developer)", "Project Manager"),
    ("Project Manager", "U.G.C. Nayanathara (also Developer)", "Day-to-day control, sprints, reporting, stakeholder liaison, application development"),
    ("Start-Up Manager", "G.G.M.P. Kusumsiri (also Developer Team Leader)", "Client liaison, mandate documentation; will lead developer team and code integration"),
    ("Quality Manager", "T.H. Samaranayake (also Developer)", "Will own CI checklist, UAT, acceptance sign-off support, and application development"),
    ("Risk Manager", "G.M.D. Fernando (also Developer)", "Will maintain risk log, monthly reviews, and application development"),
    ("Scheduling Manager", "N.D. Ranasinghagei (also Developer)", "Will maintain milestone plan, sprint calendar, and application development"),
    ("Technical Officer (school)", "School appointee", "Will handle post-go-live Firebase admin and user provisioning"),
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(table, fill: str = "1F4E79") -> None:
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, bold: bool = False, size: int = 11) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False, indent: float = 0) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        set_paragraph_text(new_para, text, bold=bold)
    if indent:
        new_para.paragraph_format.left_indent = Inches(indent)
    return new_para


def add_table_after_paragraph(doc: Document, paragraph: Paragraph, row_count: int, col_count: int):
    table = doc.add_table(rows=row_count, cols=col_count)
    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)
    paragraph._element.addnext(tbl_element)
    return table


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def find_paragraph(doc: Document, predicate) -> Paragraph | None:
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    return None


def find_section_heading(doc: Document, section_num: str, title_fragment: str) -> Paragraph | None:
    matches = []
    for p in doc.paragraphs:
        t = p.text.replace("\t", " ").strip()
        if t.startswith(section_num) and title_fragment.lower() in t.lower():
            matches.append(p)
    return matches[-1] if matches else None


def replace_block_after_heading(
    doc: Document,
    heading: Paragraph,
    stop_prefixes: tuple[str, ...],
    builder,
) -> None:
    """Remove paragraphs after heading until next section; call builder(last_para)."""
    to_remove: list[Paragraph] = []
    found = False
    for p in doc.paragraphs:
        if p is heading:
            found = True
            continue
        if not found:
            continue
        t = p.text.replace("\t", " ").strip()
        if any(t.startswith(prefix) for prefix in stop_prefixes):
            break
        if t:
            to_remove.append(p)

    anchor = heading
    for p in to_remove:
        remove_paragraph(p)

    builder(anchor)


def build_bullet_block(anchor: Paragraph, title: str, items: list[str]) -> Paragraph:
    last = anchor
    if title:
        last = insert_paragraph_after(last, title, bold=True)
    for item in items:
        last = insert_paragraph_after(last, f"• {item}", indent=0.35)
    return last


def build_numbered_block(anchor: Paragraph, items: list[str]) -> Paragraph:
    last = anchor
    for i, item in enumerate(items, 1):
        last = insert_paragraph_after(last, f"{i}. {item}", indent=0.35)
    return last


def apply_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)


def build_table(
    doc: Document,
    anchor: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    table = add_table_after_paragraph(doc, anchor, len(rows) + 1, len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Normal Table"
    except KeyError:
        pass
    apply_table_borders(table)

    for i, h in enumerate(headers):
        set_paragraph_text(table.rows[0].cells[i].paragraphs[0], h, bold=True, size=10)
    style_header_row(table)

    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, val in enumerate(row_data):
            set_paragraph_text(table.rows[r_idx].cells[c_idx].paragraphs[0], val, size=10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


def format_moscow(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: t.strip().startswith("MoSCoW prioritisation"))
    if p:
        remove_paragraph(p)

    scope_h = find_section_heading(doc, "5.3", "Project Scope")
    if not scope_h:
        return

    sub = insert_paragraph_after(scope_h, "MoSCoW prioritisation:", bold=True)
    rows = []
    for category, items in MOSCOW_ROWS:
        rows.append([category, "\n".join(f"• {x}" for x in items)])
    build_table(doc, sub, ["Priority", "Scope items"], rows, [1.3, 5.2])


def format_deliverables(doc: Document) -> None:
    def is_deliverable_block(t: str) -> bool:
        return t.strip().startswith("1. Production web application")

    p = find_paragraph(doc, is_deliverable_block)
    if not p:
        return
    h = find_section_heading(doc, "5.4", "Deliverables")
    if not h:
        return
    remove_paragraph(p)
    build_numbered_block(h, DELIVERABLES)


def format_org_structure(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: "Project Board (direction)" in t)
    if not p:
        return
    h = find_section_heading(doc, "6.1", "Project Management Team Structure")
    if not h:
        return
    remove_paragraph(p)
    rows = [[a, b, c] for a, b, c in ORG_ROWS]
    build_table(doc, h, ["Role", "Name", "Responsibility"], rows, [1.8, 1.6, 3.1])


def format_job_descriptions(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: t.strip().startswith("Project Manager (U.G.C"))
    if not p:
        return
    h = find_section_heading(doc, "6.2", "Role")
    if not h:
        h = find_section_heading(doc, "6.2", "Job")
    if not h:
        return
    remove_paragraph(p)

    jobs = [
        ("Project Manager", "U.G.C. Nayanathara (also Developer)",
         "Will plan sprints, run stand-ups, deliver demos, maintain documentation, manage releases, liaise with the Principal, and implement assigned application features."),
        ("Start-Up Manager", "G.G.M.P. Kusumsiri (also Developer Team Leader)",
         "Will maintain mandate documentation; liaise with the Principal; facilitate demos and feedback; lead the developer team; assign development tasks; review and integrate code; define application architecture; support CI builds."),
        ("Quality Manager", "T.H. Samaranayake (also Developer)",
         "Will own test strategy, CI checklist, UAT scripts, and acceptance sign-off records; implement assigned features and support testing."),
        ("Risk Manager", "G.M.D. Fernando (also Developer)",
         "Will maintain risk log, facilitate monthly risk reviews, track mitigations; implement assigned features and support testing."),
        ("Scheduling Manager", "N.D. Ranasinghagei (also Developer)",
         "Will maintain milestone plan, sprint calendar, and track slippage; implement assigned features and support testing."),
        ("Technical Officer", "School appointee",
         "Will handle post-go-live Firebase administration, user provisioning, and APK distribution."),
    ]
    rows = [[role, name, desc] for role, name, desc in jobs]
    build_table(doc, h, ["Role", "Assigned to", "Job description"], rows, [1.5, 1.5, 3.5])


def format_controls(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: t.strip().startswith("Controls:") and "Time" in t)
    if not p:
        p = find_paragraph(doc, lambda t: t.strip() == "Controls")
    if not p:
        return
    h = find_section_heading(doc, "9", "Project Controls")
    if not h:
        return
    if p.text.strip() != "Controls":
        remove_paragraph(p)

    controls = [
        ("Time", "N.D. Ranasinghagei (Scheduling Manager) maintains sprint plan and milestone tracker."),
        ("Quality", "CI pipeline and Quality Manager checklist run before each release."),
        ("Risk", "Monthly risk log review; new risks logged within 24 hours of identification."),
        ("Scope", "Change requests will require approval from the Principal (G.G. Kusumsiri), coordinated by the Start-Up Manager (G.G.M.P. Kusumsiri), when MoSCoW 'Must' items are affected."),
        ("Configuration", "Git tags for releases; production releases from main branch only."),
        ("Security", "Firestore rules publish checklist; no secrets committed to repository."),
    ]
    monitoring = [
        "GitHub Actions CI status on every push and pull request.",
        "Firebase Console usage and Authentication user counts.",
        "Application health checks and planned production support endpoint.",
    ]
    exceptions = [
        "If sprint commitment slips more than 3 days, PM reports to Scheduling Manager with a recovery plan.",
        "If a critical defect is found after go-live, use a hotfix branch; Quality Manager approves emergency release.",
        "If the Principal changes mandate scope, update the Project Brief and re-baseline this plan.",
    ]

    anchor = h
    anchor = build_bullet_block(anchor, "Controls", [f"{k} — {v}" for k, v in controls])
    anchor = build_bullet_block(anchor, "Monitoring", monitoring)
    build_bullet_block(anchor, "Exception process", exceptions)


def format_risk_log(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: "R01 |" in t or t.strip().startswith("Initial Risk Log:"))
    if p:
        remove_paragraph(p)
    h = find_section_heading(doc, "10", "Risk")
    if not h:
        return

    rows = [[a, b, c, d, e] for a, b, c, d, e in RISKS]
    build_table(
        doc,
        h,
        ["ID", "Risk description", "Impact", "Mitigation", "Owner"],
        rows,
        [0.55, 2.0, 0.7, 2.3, 1.2],
    )


def format_objectives(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: "Project delivery objectives" in t)
    if not p:
        return
    h = find_section_heading(doc, "5.1", "Objectives")
    if not h:
        return
    remove_paragraph(p)

    delivery = [
        "Complete Phase 1 pilot go-live at Delta Gemunupura College by August 2026.",
        "Deliver a production-ready web PWA and official Android APK without Play Protect warnings.",
        "Provision admin accounts and onboard at least 10 staff and 50 student/parent logins.",
        "Establish a CI pipeline (lint, unit tests, production build) that passes on every release.",
        "Design and publish Firestore security rules aligned with the planned role-based access model.",
    ]
    outcomes = [
        "Reduce time to enter and retrieve term test results by at least 50% vs manual methods.",
        "Teachers retrieve student contact details and results within 30 seconds of login.",
        "Parents gain same-day visibility of published exam results after marks are entered.",
        "System auto-calculates grades, pass/fail status, and class ranks when results are saved.",
        "Teachers restricted to assigned grades; Vice Principal will have view-only access to all modules.",
        "Maintain reliable uptime for the hosted web application.",
    ]
    anchor = build_bullet_block(h, "Project delivery objectives", delivery)
    build_bullet_block(anchor, "Project outcome objectives", outcomes)


def format_exclusions_constraints(doc: Document) -> None:
    mapping = [
        (lambda t: t.strip().startswith("• Native iOS"), "5.5", "Exclusions", [
            "Native iOS application distributed through the Apple App Store",
            "Timetable scheduling module (deferred to Phase 2)",
            "Hardware procurement (cloud hosting only)",
            "Ongoing Firebase and hosting subscription costs beyond pilot period",
            "On-site IT support contract after handover",
            "Data migration from legacy paper records beyond CSV import tooling",
            "Sinhala/Tamil full UI localisation (English UI with Sri Lankan data formats)",
        ]),
        (lambda t: t.strip().startswith("• Academic calendar"), "5.6", "Constraints", [
            "Pilot must be usable before Term 3 2026 examinations where possible",
            "Team size: 5 student project roles; G.G.M.P. Kusumsiri is Start-Up Manager and Developer Team Leader; all other team members are developers",
            "Budget: free-tier / low-cost cloud (Firebase, affordable hosting)",
            "No dedicated full-time developer staff at school post-handover",
            "Must comply with data protection expectations for student PII",
            "Must not store passwords or PIN hashes in client-accessible Firestore documents",
            "Android distribution must use signed official APK, not Chrome WebAPK",
        ]),
        (lambda t: t.strip().startswith("• School has reliable"), "5.7", "Assumptions", [
            "School has reliable internet connectivity for daily use",
            "Principal and technical officer assign time for UAT and acceptance sign-off",
            "Teachers receive short orientation on login and result entry",
            "Firebase service account and QUICK_PIN_PEPPER will be configured for production hosting",
            "GitHub Actions secrets will be available for Android APK builds",
            "Student admission numbers and staff IDs are accurate for linking logins",
            "Parents have smartphones capable of installing PWA or Android APK",
        ]),
    ]
    for pred, sec, title, items in mapping:
        p = find_paragraph(doc, pred)
        if not p:
            continue
        headings = [x for x in doc.paragraphs if sec in x.text and title in x.text]
        h = headings[-1] if headings else None
        if not h:
            continue
        remove_paragraph(p)
        build_bullet_block(h, "", items)


def format_legal_ethical(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: t.strip().startswith("Legal:") and "Personal data" in t)
    if not p:
        return
    h = find_section_heading(doc, "5.8", "Legal")
    if not h:
        return
    remove_paragraph(p)
    legal = [
        "Personal data (student NIC, parent contact, exam results) must be processed lawfully; access limited by role.",
        "Firebase and cloud hosting data processing terms apply; school remains data controller for student records.",
        "APK signing and Digital Asset Links must associate the Android app with the school domain.",
    ]
    ethical = [
        "Equitable access: parents without smartphones are not excluded (paper reports remain available).",
        "Student photos and disability notes are sensitive — visible only to authorised staff.",
        "Inactive accounts must be disabled promptly when staff leave or students transfer.",
        "Examination data must not be altered without an audit trail in Firestore records.",
    ]
    anchor = build_bullet_block(h, "Legal", legal)
    build_bullet_block(anchor, "Ethical", ethical)


def format_background(doc: Document) -> None:
    p = find_paragraph(doc, lambda t: t.startswith("CP/KOT/Delta Gemunupura College") and "manual processes" in t)
    if not p or len(p.text) < 400:
        return
    h = find_section_heading(doc, "4", "Background")
    if not h:
        return
    remove_paragraph(p)
    paras = [
        "CP/KOT/Delta Gemunupura College (Pussellawa) currently manages student, staff, examination, and inventory records through manual processes and disconnected spreadsheets. On 17 June 2026, Principal G.G. Kusumsiri issued a formal mandate requesting a comprehensive cloud-based School Management System as a pilot at the school.",
        "The mandate requires web, Android, and iOS access; core modules for students, staff, inventory, examinations, and timetables; real-time performance analytics when term test marks are entered; parental engagement through a dedicated parent portal; and strict role-based access control for principal, technical officer, vice principal, teachers, students, and parents.",
        "A Project Brief (CS11) was prepared under PRINCE2 Starting Up a Project. This Project Plan defines how the student development team will design, build, test, and hand over the Smart School Management System. The proposed solution will use Next.js, Firebase (Authentication and Firestore), and cloud-hosted hosting. Timetable management and a native iOS application are planned for a later phase.",
        "Business case: once delivered, the system will reduce administrative burden, improve accuracy of academic records, enable faster parent communication, and provide leadership with school-wide visibility — addressing a missing component of the school's digital strategy.",
    ]
    last = h
    for para in paras:
        last = insert_paragraph_after(last, para)


def format_communication_quality(doc: Document) -> None:
    comm_p = find_paragraph(doc, lambda t: "Communication Plan (summary" in t)
    if comm_p:
        h = find_section_heading(doc, "7", "Communication")
        if h:
            remove_paragraph(comm_p)
            rows = [
                ["Weekly sprint stand-up", "Development team", "30 minutes", "Scheduling Manager (minutes)"],
                ["Fortnightly demo", "Principal (G.G. Kusumsiri) and Start-Up Manager", "45 minutes", "U.G.C. Nayanathara (sprint increments)"],
                ["Sprint status report", "Project Manager", "Every Friday", "U.G.C. Nayanathara"],
                ["Risk register review", "Risk Manager + team", "Monthly", "G.M.D. Fernando"],
                ["Quality release checklist", "Quality Manager", "Before each release", "T.H. Samaranayake"],
                ["Escalation (>48h blocker)", "Principal (G.G. Kusumsiri) via Start-Up Manager", "As needed", "G.G.M.P. Kusumsiri"],
            ]
            build_table(doc, h, ["Activity", "Audience", "Frequency", "Owner"], rows, [1.8, 1.8, 1.0, 1.9])

    qual_p = find_paragraph(doc, lambda t: "Project Quality Plan (summary" in t)
    if qual_p:
        h = find_section_heading(doc, "8", "Quality")
        if h:
            remove_paragraph(qual_p)
            rows = [
                ["Code quality", "TypeScript + ESLint clean on CI", "Every PR", "G.G.M.P. Kusumsiri"],
                ["Unit tests", "Unit tests for role-based access, grading, and PIN validation", "Every PR", "Developer team (Leader: G.G.M.P. Kusumsiri)"],
                ["Security", "Firestore rules aligned with planned role matrix", "Each release", "G.G.M.P. Kusumsiri"],
                ["CI build", "Automated lint, test, and build in GitHub Actions", "Every PR", "Developer team (Leader: G.G.M.P. Kusumsiri)"],
                ["Code review", "Review pull requests before merge to main", "Every PR", "T.H. Samaranayake"],
                ["UAT", "Manual checklist per module before pilot", "Pre go-live", "T.H. Samaranayake"],
                ["Acceptance", "Acceptance criteria checklist and sign-off coordination", "Phase 1 end", "T.H. Samaranayake"],
            ]
            build_table(doc, h, ["Area", "Standard / activity", "When", "Owner"], rows, [1.1, 2.6, 1.0, 1.8])


def main() -> None:
    doc = Document(DOC_PATH)
    format_background(doc)
    format_objectives(doc)
    format_moscow(doc)
    format_deliverables(doc)
    format_exclusions_constraints(doc)
    format_legal_ethical(doc)
    format_org_structure(doc)
    format_job_descriptions(doc)
    format_communication_quality(doc)
    format_controls(doc)
    format_risk_log(doc)
    doc.save(DOC_PATH)
    print(f"Formatted: {DOC_PATH}")


if __name__ == "__main__":
    main()
