"""Fill the PRINCE2 Project Plan template for Smart School Management System."""

from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

TEMPLATE = Path(__file__).parent / "Project plan template.docx"
OUTPUT = Path(__file__).parent / "Project Plan - Delta Gemunupura DBMS.docx"

TODAY = date.today().strftime("%d %B %Y")
NEXT_REVIEW = "17 July 2026"

# Match placeholders by unique substring (handles smart quotes in template)
PLACEHOLDER_CONTENT: list[tuple[str, str]] = [
    (
        "Explaining the context of the project",
        (
            "CP/KOT/Delta Gemunupura College (Pussellawa) currently manages student, staff, examination, "
            "and inventory records through manual processes and disconnected spreadsheets. On 17 June 2026, "
            "Principal G.G. Kusumsiri issued a formal mandate requesting a comprehensive cloud-based School "
            "Management System as a pilot at the school.\n\n"
            "The mandate requires web, Android, and iOS access; core modules for students, staff, inventory, "
            "examinations, and timetables; real-time performance analytics when term test marks are entered; "
            "parental engagement through a dedicated parent portal; and strict role-based access control for "
            "principal, technical officer, vice principal, teachers, students, and parents.\n\n"
            "A Project Brief (CS11) was prepared under PRINCE2 Starting Up a Project. This Project Plan defines "
            "how the student development team will design, build, test, and hand over the Smart School Management "
            "System. The proposed solution will use Next.js, Firebase (Authentication and Firestore), and "
            "cloud-hosted hosting. Timetable management and a native iOS application are planned for a later phase.\n\n"
            "Business case: once delivered, the system will reduce administrative burden, improve accuracy of "
            "academic records, enable faster parent communication, and provide leadership with school-wide "
            "visibility — addressing a missing component of the school's digital strategy."
        ),
    ),
    (
        "Specifically what is required to be achieved",
        (
            "Project delivery objectives:\n"
            "• Complete Phase 1 pilot go-live at Delta Gemunupura College by August 2026.\n"
            "• Deliver a production-ready web PWA and official Android APK installable without Play Protect warnings.\n"
            "• Provision principal and technical officer admin accounts; onboard at least 10 staff and 50 student/parent logins for the pilot.\n"
            "• Establish a CI pipeline (lint, unit tests, production build) that passes on every release.\n"
            "• Design and publish Firestore security rules aligned with the planned role-based access model.\n\n"
            "Project outcome objectives (measurable after go-live):\n"
            "• Reduce time to enter and retrieve term test results by at least 50% compared to manual methods.\n"
            "• Enable teachers to retrieve a student's contact details and results within 30 seconds of login.\n"
            "• Provide parents same-day visibility of published exam results after marks are entered.\n"
            "• Auto-calculate subject grades, pass/fail status, and class ranks when results are saved.\n"
            "• Restrict each teacher to students in assigned grades only; vice principal will have view-only access to all modules.\n"
            "• Maintain reliable uptime for the hosted web application."
        ),
    ),
    (
        "how the work of the project is going to be approached",
        (
            "The project will be governed using PRINCE2 for documentation, roles, controls, and risk management. "
            "Delivery will follow an Agile iterative approach: two-week sprints, demonstrable software at the end of each "
            "sprint, sprint reviews with the Project Manager and Quality Manager, and backlog prioritisation "
            "aligned to the principal's mandate.\n\n"
            "Proposed technical approach: Next.js App Router, Firebase Authentication and Firestore, Tailwind CSS, "
            "role-based access controls mirrored in Firestore security rules, GitHub Actions CI/CD, and a Bubblewrap "
            "Trusted Web Activity for Android distribution. Application development will be carried out by the developer team, "
            "led by G.G.M.P. Kusumsiri (Developer Team Leader), with U.G.C. Nayanathara, T.H. Samaranayake, G.M.D. Fernando, and N.D. Ranasinghagei as developers."
        ),
    ),
    (
        "scoping diagram",
        (
            "MoSCoW prioritisation:\n\n"
            "Must have:\n"
            "• Student Management (CRUD, search, CSV import, PDF export)\n"
            "• Staff Management (CRUD, Google Form CSV/XLSX import)\n"
            "• Examination Management (create exams, enter results, auto-grading, ranks, report export)\n"
            "• Inventory Management (asset register CRUD)\n"
            "• User Management (role-based logins: principal, technical officer, vice principal, teacher, student, parent)\n"
            "• Web PWA with install support\n"
            "• Android official APK (TWA)\n"
            "• Real-time performance analytics on data entry\n"
            "• Parent portal (view child's progress and exam results)\n\n"
            "Should have:\n"
            "• Quick PIN login for mobile users\n"
            "• Biometric unlock on installed PWA\n"
            "• WhatsApp integration for sharing student credentials\n"
            "• Examination information reports for leadership\n\n"
            "Could have:\n"
            "• Timetable Management module\n"
            "• Teacher performance percentage dashboards\n"
            "• Bulk SMS notifications to parents\n\n"
            "Won't have (Phase 1):\n"
            "• Native iOS App Store application (iOS served via Safari PWA install only)\n"
            "• Full offline database synchronisation\n"
            "• Integration with government EMIS or national exam systems\n"
            "• Fee/billing or payroll modules"
        ),
    ),
    (
        "expected and required Deliverables",
        (
            "1. Production web application (PWA) for school use\n"
            "2. Android APK (Trusted Web Activity) with install guide\n"
            "3. Firebase project with Authentication, Firestore, security rules, and composite indexes\n"
            "4. Source code repository with CI workflow\n"
            "5. User documentation: system guide, Android install guide, and iOS install instructions\n"
            "6. Sample CSV templates for student, staff, exam results, and inventory import\n"
            "7. Admin setup utilities for creating the first principal account and provisioning student logins\n"
            "8. Project documentation: Project Mandate (principal letter), Project Brief, and this Project Plan\n"
            "9. Role-based dashboards and module screens for all mandated user types\n"
            "10. PDF export for student/staff profiles and examination reports"
        ),
    ),
    (
        "What is not included in the project",
        (
            "• Native iOS application distributed through the Apple App Store\n"
            "• Timetable scheduling module (planned for Phase 2)\n"
            "• Hardware procurement (tablets, servers — cloud hosting only)\n"
            "• Ongoing Firebase and hosting subscription costs beyond pilot period\n"
            "• On-site IT support contract after handover\n"
            "• Data migration from legacy paper records beyond CSV import tooling\n"
            "• Sinhala/Tamil full UI localisation (English UI with Sri Lankan data formats)"
        ),
    ),
    (
        "Restrictions on time, resources, funding",
        (
            "• Academic calendar: pilot must be usable before Term 3 2026 examinations where possible\n"
            "• Team size: 5 student project roles; G.G.M.P. Kusumsiri is Start-Up Manager and Developer Team Leader; all other team members are developers\n"
            "• Budget: free-tier / low-cost cloud (Firebase Spark/Blaze and affordable hosting)\n"
            "• No dedicated full-time developer staff at the school — maintenance by technical officer post-handover\n"
            "• Must comply with school and national data protection expectations for student PII\n"
            "• Must not store passwords or PIN hashes in client-accessible Firestore documents\n"
            "• Android distribution must use signed official APK, not Chrome WebAPK (Play Protect requirement)"
        ),
    ),
    (
        "Similar to Constraints but more",
        (
            "• School has reliable internet connectivity for daily use\n"
            "• Principal and technical officer will assign time for UAT and user acceptance sign-off\n"
            "• Teachers will receive a short orientation on login and result entry\n"
            "• Firebase service account and QUICK_PIN_PEPPER secrets will be configured for production hosting\n"
            "• GitHub Actions secrets will be available for Android APK builds\n"
            "• Student admission numbers and staff IDs from existing school registers are accurate for linking logins\n"
            "• Parents have smartphones capable of installing PWA or Android APK"
        ),
    ),
    (
        "legal and/or ethical issues",
        (
            "Legal:\n"
            "• Personal data (student NIC, parent contact, exam results) must be processed lawfully; access limited by role\n"
            "• Firebase and cloud hosting data processing terms apply; school remains data controller for student records\n"
            "• APK signing and Digital Asset Links must correctly associate the Android app with the school domain\n\n"
            "Ethical:\n"
            "• Equitable access: ensure parents without smartphones are not excluded (paper reports remain available)\n"
            "• Student photos and disability notes are sensitive — visible only to authorised staff\n"
            "• Inactive accounts must be disabled promptly when staff leave or students transfer\n"
            "• Examination data must not be altered without an audit trail in Firestore records"
        ),
    ),
    (
        "Communication Plan Template",
        (
            "Communication Plan (summary — see project files for full detail):\n\n"
            "Stakeholders: Principal (client), technical officer, vice principal, teachers, parents, development team.\n\n"
            "Channels:\n"
            "• Weekly sprint stand-up (team, 30 min) — Scheduling Manager maintains minutes\n"
            "• Fortnightly demo to Principal (G.G. Kusumsiri), coordinated by Start-Up Manager (G.G.M.P. Kusumsiri) — U.G.C. Nayanathara will present sprint increments\n"
            "• GitHub repository for code, issues, and pull request reviews\n"
            "• WhatsApp group for urgent school-side queries (principal contact: 0776987325)\n"
            "• Email for formal approvals and document distribution\n\n"
            "Reporting:\n"
            "• Sprint status report to Project Manager every Friday\n"
            "• Risk register reviewed monthly by Risk Manager (G.M.D. Fernando)\n"
            "• Quality checklist run before each release by Quality Manager (T.H. Samaranayake)\n\n"
            "Escalation: blockers > 48 hours escalated to Principal (G.G. Kusumsiri) via Start-Up Manager (G.G.M.P. Kusumsiri)."
        ),
    ),
    (
        "Project Quality Plan Template",
        (
            "Project Quality Plan (summary):\n\n"
            "Quality objectives: the delivered system must be effective, easy to understand, secure, and maintainable.\n\n"
            "Planned standards:\n"
            "• TypeScript strict typing; ESLint clean on CI\n"
            "• Unit tests for role-based access, exam grading, and Quick PIN validation\n"
            "• Firestore security rules aligned with the planned role matrix\n"
            "• Mobile-first responsive UI; PWA installable on Android and iOS Safari\n\n"
            "Planned quality activities:\n"
            "• Code review on all pull requests — developer team (Leader: G.G.M.P. Kusumsiri); quality review by T.H. Samaranayake\n"
            "• Automated lint, test, and build checks in GitHub Actions CI — developer team (Leader: G.G.M.P. Kusumsiri)\n"
            "• Manual UAT checklist per module before pilot go-live — T.H. Samaranayake\n"
            "• Acceptance criteria checklist and sign-off coordination — T.H. Samaranayake\n\n"
            "Quality roles: T.H. Samaranayake (Quality Manager) will lead quality review and UAT; G.G.M.P. Kusumsiri (Developer Team Leader) will lead development quality and CI."
        ),
    ),
    (
        "how control is to be exercised within the project",
        (
            "Controls:\n"
            "1. Time — N.D. Ranasinghagei (Scheduling Manager) maintains sprint plan and milestone tracker\n"
            "2. Quality — CI pipeline + Quality Manager checklist before each release\n"
            "3. Risk — monthly risk log review; new risks logged within 24 hours of identification\n"
            "4. Scope — change requests will require approval from the Principal (client, G.G. Kusumsiri), coordinated by the Start-Up Manager, if MoSCoW 'Must' items are affected\n"
            "5. Configuration — Git tags for releases; production releases from main branch only\n"
            "6. Security — Firestore rules publish checklist; no secrets in repository\n\n"
            "Monitoring:\n"
            "• GitHub Actions CI status on every push/PR (once CI is established)\n"
            "• Firebase Console usage and Auth user counts after go-live\n"
            "• Planned application health-check endpoint for production support\n\n"
            "Exception process:\n"
            "• If sprint commitment slips > 3 days, PM will report to Scheduling Manager and document a recovery plan\n"
            "• If a critical defect is found after go-live, use a hotfix branch; Quality Manager approves emergency release\n"
            "• If mandate scope changes from the Principal, update the Project Brief and re-baseline this plan"
        ),
    ),
    (
        "risk analysis and risk management activities",
        (
            "Initial Risk Log:\n\n"
            "R01 | Internet dependency — no offline database | High | Plan PWA offline page; communicate network requirement to users | G.M.D. Fernando\n"
            "R02 | Firebase cost overrun at scale | Medium | Monitor usage; optimise queries; publish composite indexes | G.M.D. Fernando\n"
            "R03 | Teachers resist adopting new system | Medium | Training session; Quick PIN login; parallel paper records during pilot | G.M.D. Fernando\n"
            "R04 | Incorrect RBAC exposes student data | High | Design Firestore rules and access-control tests; verify VP view-only | G.M.D. Fernando\n"
            "R05 | iOS users expect native App Store app | Low | Provide iOS install guide; document exclusion in Project Brief | G.M.D. Fernando\n"
            "R06 | Android Play Protect warnings (WebAPK) | Medium | Plan official signed Android APK; document install steps for IT | G.M.D. Fernando\n"
            "R07 | Service account / PIN pepper misconfigured for production | High | Use production environment setup checklist; health-check endpoint | G.M.D. Fernando\n"
            "R08 | Timetable module not ready for pilot | Medium | Defer to Phase 2; agree scope with Principal | G.M.D. Fernando\n"
            "R09 | Student login linking errors (wrong admission number) | Medium | Plan bulk provision process; admin verification before go-live | G.M.D. Fernando\n"
            "R10 | Academic calendar pressure | Medium | Prioritise exam module in sprint plan; agile reprioritisation each sprint | G.M.D. Fernando"
        ),
    ),
]

ORG_STRUCTURE = (
    "Project Board (direction):\n"
    "• Executive / Client: G.G. Kusumsiri (Principal, CP/KOT/Delta Gemunupura College)\n"
    "• Senior User: Vice Principal (school operations representative)\n"
    "• Senior Supplier: U.G.C. Nayanathara (Project Manager, also Developer)\n\n"
    "Project team:\n"
    "• Project Manager: U.G.C. Nayanathara (also Developer)\n"
    "• Start-Up Manager: G.G.M.P. Kusumsiri (also Developer Team Leader)\n"
    "• Quality Manager: T.H. Samaranayake (also Developer)\n"
    "• Risk Manager: G.M.D. Fernando (also Developer)\n"
    "• Scheduling Manager: N.D. Ranasinghagei (also Developer)\n\n"
    "Lines of authority: all developers report to G.G.M.P. Kusumsiri (Developer Team Leader); the Developer Team Leader reports to the Project Manager. "
    "Scope and acceptance decisions are approved by the Principal (client). The Start-Up Manager coordinates between the team and the Principal."
)

JOB_DESCRIPTIONS = (
    "Project Manager (U.G.C. Nayanathara, also Developer): Will plan sprints, run stand-ups, deliver demos, maintain documentation, manage releases, "
    "liaise with the Principal, and implement assigned application features.\n\n"
    "Start-Up Manager (G.G.M.P. Kusumsiri, also Developer Team Leader): Will maintain mandate documentation; liaise with the Principal (client); "
    "facilitate demos and feedback; escalate scope questions to the Principal; lead the developer team; assign development tasks; "
    "review and integrate code; define the application architecture; and support CI builds.\n\n"
    "Quality Manager (T.H. Samaranayake, also Developer): Will own test strategy, CI checklist, UAT scripts, and acceptance sign-off records; "
    "implement assigned application features and support testing.\n\n"
    "Risk Manager (G.M.D. Fernando, also Developer): Will maintain the risk log, facilitate monthly risk reviews, track mitigations; "
    "implement assigned application features and support testing.\n\n"
    "Scheduling Manager (N.D. Ranasinghagei, also Developer): Will maintain the milestone plan, sprint calendar, and track slippage; "
    "implement assigned application features and support testing.\n\n"
    "Technical Officer (school): Will handle post-go-live Firebase administration, user provisioning, and APK distribution after handover."
)

COVER_VALUES = {
    "Project:": "Smart School Management System",
    "Release:": "Phase 1: Pilot",
    "Date:": TODAY,
    "PM Approach:": "PRINCE2 governance with Agile (2-week sprints)",
    "Author:": "G.G.M.P. Kusumsiri",
    "Owner:": "G.G.M.P. Kusumsiri (Start-Up Manager)",
    "Client:": "G.G. Kusumsiri,\nPrincipal (SLPS-1),\nCP/KOT/Delta Gemunupura College,\nPussellawa",
    "Document Ref:": "DGC-DBMS-PP-001",
    "Version No:": "1.0",
}


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def match_placeholder(text: str) -> str | None:
    if not text.strip().startswith("["):
        return None
    for key, content in PLACEHOLDER_CONTENT:
        if key.lower() in text.lower():
            return content
    return None


def replace_placeholders_in_paragraph(paragraph: Paragraph) -> bool:
    content = match_placeholder(paragraph.text)
    if content is None:
        return False
    paragraph.text = content
    return True


def replace_all_placeholders(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph)


def insert_after_heading(doc: Document, heading_substring: str, content: str) -> None:
    for i, paragraph in enumerate(doc.paragraphs):
        if heading_substring in paragraph.text and not paragraph.text.strip().startswith("["):
            for j in range(i + 1, len(doc.paragraphs)):
                nxt = doc.paragraphs[j].text.strip()
                if nxt.startswith("[") or match_placeholder(doc.paragraphs[j].text):
                    doc.paragraphs[j].text = content
                    return
                if any(
                    h in nxt
                    for h in (
                        "6.2",
                        "7Communication",
                        "7\tCommunication",
                        "Individual Role",
                        "Job Descriptions",
                    )
                ):
                    new_p = paragraph._element
                    # insert new paragraph before next heading
                    from docx.oxml import OxmlElement

                    p = OxmlElement("w:p")
                    r = OxmlElement("w:r")
                    t = OxmlElement("w:t")
                    t.text = content
                    r.append(t)
                    p.append(r)
                    doc.paragraphs[j]._element.addprevious(p)
                    return
            return


def fill_metadata_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        t = paragraph.text.strip()
        if t == "The source of the document will be found on the project's PC in location":
            paragraph.text = (
                "The source of the document will be found on the project's PC in location "
                r"C:\Users\User\Documents\GitHub\delta-gemunupura-dbms-2026\agile project"
            )
        elif t == "Date of this revision:":
            paragraph.text = f"Date of this revision: {TODAY}"
        elif t == "Date of Next revision:":
            paragraph.text = f"Date of Next revision: {NEXT_REVIEW}"


def fill_tables(doc: Document) -> None:
    if not doc.tables:
        return

    # Cover table
    for row in doc.tables[0].rows:
        label = row.cells[0].text.strip()
        if label in COVER_VALUES and len(row.cells) > 1:
            row.cells[1].text = COVER_VALUES[label]

    # Revision table
    if len(doc.tables) > 1:
        row = doc.tables[1].rows[1]
        row.cells[2].text = "First issue"

    # Approvals
    if len(doc.tables) > 2:
        approvals = [
            ("G.G. Kusumsiri", "Principal / Client", TODAY, "1.0"),
            ("U.G.C. Nayanathara", "Project Manager", TODAY, "1.0"),
            ("T.H. Samaranayake", "Quality Manager", TODAY, "1.0"),
        ]
        table = doc.tables[2]
        while len(table.rows) < len(approvals) + 1:
            table.add_row()
        for idx, (name, title, dt, ver) in enumerate(approvals, start=1):
            r = table.rows[idx]
            r.cells[0].text = name
            r.cells[2].text = title
            r.cells[3].text = dt
            r.cells[4].text = ver

    # Distribution
    if len(doc.tables) > 3:
        dist = [
            ("G.G. Kusumsiri", "Principal / Client", TODAY, "1.0"),
            ("G.G.M.P. Kusumsiri", "Start-Up Manager (also Developer Team Leader)", TODAY, "1.0"),
            ("U.G.C. Nayanathara", "Project Manager (also Developer)", TODAY, "1.0"),
            ("T.H. Samaranayake", "Quality Manager (also Developer)", TODAY, "1.0"),
            ("G.M.D. Fernando", "Risk Manager (also Developer)", TODAY, "1.0"),
            ("N.D. Ranasinghagei", "Scheduling Manager (also Developer)", TODAY, "1.0"),
        ]
        table = doc.tables[3]
        while len(table.rows) < len(dist) + 1:
            table.add_row()
        for idx, (name, title, dt, ver) in enumerate(dist, start=1):
            r = table.rows[idx]
            r.cells[0].text = name
            r.cells[1].text = title
            r.cells[2].text = dt
            r.cells[3].text = ver


def remove_template_instructions(doc: Document) -> None:
    remove_phrases = (
        "how to use this template",
        "once your plan is completed",
        "does the document correctly",
        "does it show a viable",
        "is the project organisation structure complete",
        "have all the roles been considered",
        "is the project organisation structure backed up",
        "are the relationships and lines of authority",
        "do the controls satisfy",
        "does it clearly show a control",
        "is it clear who will administer",
        "project documentation",
    )
    for paragraph in list(doc.paragraphs):
        t = normalize(paragraph.text)
        if any(p in t for p in remove_phrases):
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)


def main() -> None:
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    remove_template_instructions(doc)
    replace_all_placeholders(doc)
    fill_metadata_paragraphs(doc)
    fill_tables(doc)

    # Section 6 content after second occurrence of headings (body, not TOC)
    seen_61 = 0
    seen_62 = 0
    for i, paragraph in enumerate(doc.paragraphs):
        t = paragraph.text.strip()
        if "6.1" in t and "Project Management Team Structure" in t:
            seen_61 += 1
            if seen_61 == 2:
                for j in range(i + 1, len(doc.paragraphs)):
                    if doc.paragraphs[j].text.strip().startswith("6.2"):
                        doc.paragraphs[j]._element.addprevious(_new_paragraph_element(doc, ORG_STRUCTURE))
                        break
        if ("6.2" in t) and ("Job Descriptions" in t or "Individual Role" in t):
            seen_62 += 1
            if seen_62 == 2:
                for j in range(i + 1, len(doc.paragraphs)):
                    if doc.paragraphs[j].text.strip().startswith("7"):
                        doc.paragraphs[j]._element.addprevious(_new_paragraph_element(doc, JOB_DESCRIPTIONS))
                        break

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")

    # Apply table/bullet formatting
    import subprocess
    import sys
    fmt = Path(__file__).parent / "format-project-plan.py"
    subprocess.run([sys.executable, str(fmt)], check=True)


def _new_paragraph_element(doc: Document, text: str):
    from docx.oxml import OxmlElement

    p = OxmlElement("w:p")
    for line in text.split("\n"):
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = line
        r.append(t)
        p.append(r)
        if line != text.split("\n")[-1]:
            br = OxmlElement("w:br")
            p.append(br)
    return p


if __name__ == "__main__":
    main()
